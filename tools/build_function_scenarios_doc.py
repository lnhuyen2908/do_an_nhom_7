from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_FILE = OUT_DIR / "Kich_ban_chuc_nang_Web_Quan_ly_Trung_tam_Tieng_Anh_Nhom_7.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B4D"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F5F8FC"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"
GREEN = "1E6B45"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    # Keep the repeating header with at least the first body row so a table
    # never leaves only its header at the bottom of a page.
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_run_font(run, name="Calibri", size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_table_text_style(table, body_size=9.4) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=body_size if row_index else 9.2,
                        bold=True if row_index == 0 else None,
                        color=INK,
                    )


def add_table(doc, headers: list[str], rows: Iterable[Iterable[str]], widths_dxa: list[int],
              header_fill=LIGHT_BLUE, body_size=9.4):
    rows = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    for idx, label in enumerate(headers):
        header.cells[idx].text = label
        set_cell_shading(header.cells[idx], header_fill)
    repeat_table_header(header)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            if row_index % 2 == 1:
                set_cell_shading(cells[idx], LIGHTER_BLUE)
    set_table_geometry(table, widths_dxa)
    set_table_text_style(table, body_size=body_size)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(3)
    after.paragraph_format.line_spacing = Pt(1)
    tiny = after.add_run(" ")
    set_run_font(tiny, size=1, color=WHITE)
    return table


def add_label_paragraph(doc, label: str, text: str, color=INK, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=11, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_callout(doc, label: str, text: str, fill=LIGHT_BLUE, color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    # Treat the single callout row as its semantic header so accessibility
    # auditors do not interpret it as an unlabeled data table.
    repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=10.2, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=INK)
    set_table_geometry(table, [9360])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    after.paragraph_format.line_spacing = Pt(1)
    tiny = after.add_run(" ")
    set_run_font(tiny, size=1, color=WHITE)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 28, INK, 0, 8),
        ("Subtitle", 13.5, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.1


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("NHÓM 7 | WEB QUẢN LÝ TRUNG TÂM TIẾNG ANH")
    set_run_font(r, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("Kịch bản chức năng | Trang ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_number(p)


def metadata_rows(s):
    return [
        ("Tác nhân", s["actors"]),
        ("Khởi tạo", s["initiator"]),
        ("Điểm bắt đầu", s["entry"]),
        ("Tiền điều kiện", s["prereq"]),
        ("Kết thúc thành công", s["success"]),
        ("Giới hạn", s["limits"]),
    ]


SCENARIOS = [
    {
        "id": "F01",
        "title": "Trang chủ, tiện ích chung và liên hệ",
        "actors": "Khách; Admin; Staff; Teacher; Student; dịch vụ từ điển MyMemory; Zalo.",
        "initiator": "Bất kỳ người dùng nào mở trang chủ hoặc dùng tiện ích nổi.",
        "entry": "Home/Index; thanh điều hướng; nút Từ điển; nút Liên hệ Zalo.",
        "prereq": "Ứng dụng hoạt động. Đăng nhập chỉ cần khi xem số thông báo chưa đọc.",
        "success": "Hiển thị số liệu tổng quan, 3 khóa học nổi bật; mở Zalo hoặc trả nghĩa Anh - Việt.",
        "limits": "Từ điển chỉ nhận chữ tiếng Anh; tra sau 400 ms; API ngoài có thể lỗi; không ghi dữ liệu nghiệp vụ.",
        "steps": [
            ("1", "Người dùng", "Mở trang chủ.", "HomeController.Index đếm Students, Teachers, Courses, CourseClasses; lấy 3 Course theo Code."),
            ("2", "Hệ thống", "Hiển thị nội dung theo trạng thái đăng nhập.", "Nếu đã đăng nhập, đếm Notification chưa đọc của đúng UserAccount và hiện badge."),
            ("3", "Người dùng", "Bấm Xem khóa học hoặc Chi tiết.", "Đi tới Courses/Index hoặc Courses/Details."),
            ("4", "Người dùng", "Bấm Liên hệ Zalo.", "Mở https://zalo.me/0865859102 ở tab mới; hệ thống không lưu lịch sử liên hệ."),
            ("5", "Người dùng", "Mở Từ điển và nhập từ.", "site.js chuẩn hóa từ; tra từ điển cục bộ/cache, nếu thiếu thì gọi MyMemory en|vi và hiển thị kết quả."),
        ],
        "interaction": "Không có bước chờ tác nhân nội bộ. Khi demo, có thể dùng một trình duyệt khách để minh họa trang công khai; đăng nhập tài khoản bất kỳ để thấy badge thông báo.",
        "constraints": [
            ("R1", "Từ trống", "Hiện 'Nhập một từ để tra nhanh'."),
            ("R2", "Có ký tự ngoài a-z, khoảng trắng hoặc dấu gạch nối", "Hiện thông báo chỉ hỗ trợ từ tiếng Anh."),
            ("R3", "API ngoài không kết nối", "Hiện lỗi kết nối và cho phép thử lại."),
            ("R4", "Toast và xác nhận form", "Thông báo tự ẩn sau khoảng 2 giây; form có data-confirm sẽ hỏi xác nhận trước khi gửi."),
        ],
        "alternatives": [
            ("Không đăng nhập", "Badge thông báo bằng 0 và hiển thị nút Đăng ký/Đăng nhập."),
            ("Admin/Staff", "Logo và đăng nhập điều hướng tới Dashboard; trang vẫn có liên kết xem khu công khai."),
        ],
        "result": "Chỉ đọc dữ liệu. Kết quả hỗ trợ điều hướng sang tra cứu khóa học, lịch khai giảng, đăng ký tài khoản và các khu vực theo vai trò.",
        "source": "HomeController.Index; Views/Home/Index.cshtml; Views/Shared/_Layout.cshtml; wwwroot/js/site.js",
    },
    {
        "id": "F02",
        "title": "Tra cứu danh sách và chi tiết khóa học",
        "actors": "Khách, Student, Teacher, Admin, Staff.",
        "initiator": "Người dùng chọn Khóa học trên thanh điều hướng.",
        "entry": "Courses/Index; từ trang chủ hoặc lịch khai giảng mở Courses/Details/{id}.",
        "prereq": "Không cần đăng nhập để xem. Cần có Course; chi tiết lớp lấy từ CourseClasses.",
        "success": "Tìm đúng khóa học, xem nội dung, học phí, giáo viên, lớp, trạng thái và sĩ số đã duyệt.",
        "limits": "6 khóa học/trang; ngày khai giảng lọc theo ngày chính xác; sĩ số chỉ đếm Enrollment=Approved.",
        "steps": [
            ("1", "Người dùng", "Nhập từ khóa, chọn trình độ hoặc ngày khai giảng và bấm Tìm kiếm.", "CoursesController.Index trim dữ liệu, lọc Code/Name, Level và lớp có StartDate đúng ngày."),
            ("2", "Hệ thống", "Hiển thị kết quả và phân trang.", "Đếm tổng, giới hạn page trong khoảng hợp lệ, sắp theo Code hoặc ngày lớp khi có bộ lọc ngày."),
            ("3", "Người dùng", "Bấm Chi tiết.", "CoursesController.Details nạp Course, CourseClasses, Teacher; tính số chỗ đã dùng."),
            ("4", "Hệ thống", "Hiển thị mục tiêu, chương trình, học phí và danh sách lớp.", "Tính EffectiveStatus/CanRegister từ ngày và trạng thái lớp; với Student còn xác định IsSaved và HasEnrollment."),
        ],
        "interaction": "Đây là luồng đọc một phía. Từ chi tiết, Student có thể chuyển sang F07 (lưu khóa) hoặc F08 (đăng ký lớp); Admin/Staff chuyển sang F12 (sửa khóa học).",
        "constraints": [
            ("R1", "keyword rỗng hoặc chỉ gồm dấu + khi có tham số", "Hiện 'Vui lòng nhập từ khóa hợp lệ'."),
            ("R2", "Không có kết quả", "Hiện 'Không tìm thấy khóa học phù hợp'."),
            ("R3", "id thiếu hoặc không tồn tại", "Trả 404 NotFound."),
            ("R4", "Người dùng không phải Student", "Không hiện nút lưu/đăng ký; Admin/Staff thấy nút chỉnh sửa."),
        ],
        "alternatives": [("Khóa chưa có lớp", "Hiện trạng thái chưa có lớp khai giảng; không có form đăng ký."), ("Lớp đầy/khóa/đóng", "Nút đăng ký bị vô hiệu hóa và hiện lớp không nhận đăng ký.")],
        "result": "Không thay đổi dữ liệu; là đầu vào cho SavedCourse, Enrollment và Course CRUD.",
        "source": "CoursesController.Index, Details; Models/CourseClass.EffectiveStatus; Views/Courses/Index.cshtml, Details.cshtml",
    },
    {
        "id": "F03",
        "title": "Tra cứu lịch khai giảng",
        "actors": "Khách, Student, Admin, Staff và Teacher.",
        "initiator": "Người dùng bấm Lịch khai giảng.",
        "entry": "CourseClasses/Schedule.",
        "prereq": "Không cần đăng nhập. Với Teacher phải có claim TeacherId hợp lệ để xem lịch của mình.",
        "success": "Danh sách lớp đúng bộ lọc, kèm khóa học, giáo viên, lịch, phòng, ngày và trạng thái hiệu lực.",
        "limits": "10 lớp/trang; Online được nhận biết bằng Room chứa chữ 'Online'; ngày lọc theo StartDate chính xác.",
        "steps": [
            ("1", "Người dùng", "Nhập mã/tên khóa, chọn Online/Offline hoặc ngày khai giảng.", "CourseClassesController.Schedule lọc Code, Course.Code, Course.Name, Room và StartDate."),
            ("2", "Hệ thống", "Kiểm tra vai trò Teacher.", "Nếu là Teacher, chỉ giữ lớp có TeacherId trùng claim và gắn cờ lịch giáo viên."),
            ("3", "Hệ thống", "Sắp xếp và phân trang.", "Sắp StartDate giảm dần, rồi Code; giới hạn trang trong miền hợp lệ."),
            ("4", "Người dùng", "Bấm Chi tiết tại một lớp.", "Chuyển sang chi tiết Course để xem toàn bộ lớp và có thể đăng ký."),
        ],
        "interaction": "Không có thông báo hoặc chờ duyệt. Khi demo Teacher, cùng URL tự động biến thành lịch chỉ của giáo viên đang đăng nhập.",
        "constraints": [("R1", "Teacher thiếu TeacherId", "Forbid 403."), ("R2", "Không có kết quả", "Hiện bảng rỗng với thông báo phù hợp."), ("R3", "Trạng thái", "Upcoming/Open được suy ra theo ngày, trừ Locked/Closed được giữ nguyên.")],
        "alternatives": [("Khách/Student/Admin/Staff", "Xem toàn bộ lịch theo bộ lọc."), ("Teacher", "Chỉ xem lịch do chính mình phụ trách.")],
        "result": "Không ghi dữ liệu; dẫn sang chi tiết khóa học hoặc hỗ trợ Teacher kiểm tra lịch.",
        "source": "CourseClassesController.Schedule; CourseClass.EffectiveStatus; Views/CourseClasses/Schedule.cshtml",
    },
    {
        "id": "F04",
        "title": "Đăng ký tài khoản học viên và xác minh OTP",
        "actors": "Khách, EmailSender/Gmail SMTP, hệ thống.",
        "initiator": "Khách bấm Đăng ký và gửi form.",
        "entry": "Auth/Register -> Auth/VerifyOtp.",
        "prereq": "Chưa đăng nhập; vai trò Student tồn tại; email/điện thoại/tên đăng nhập chưa được dùng.",
        "success": "Tạo một Student và một UserAccount Student liên kết 1-1; OTP session bị xóa; chuyển sang Login.",
        "limits": "Mã OTP 6 chữ số, hết hạn 10 phút; session idle 20 phút; không có giới hạn số lần nhập/gửi lại trong code.",
        "steps": [
            ("1", "Khách", "Nhập họ tên, username, email, điện thoại, ngày sinh, địa chỉ, mật khẩu và xác nhận.", "AuthController.Register chuẩn hóa username chữ thường, trim dữ liệu và giữ lại dữ liệu form qua ViewBag khi lỗi."),
            ("2", "Hệ thống", "Kiểm tra bắt buộc, định dạng, độ tuổi và trùng lặp.", "Kiểm tra username >=3, password >=6, xác nhận khớp, email/phone hợp lệ, tuổi 5-100; dò UserAccount/Student."),
            ("3", "Hệ thống", "Sinh OTP và lưu đăng ký tạm.", "Tạo PendingRegistration trong session với thời hạn 10 phút; chưa ghi Student/UserAccount."),
            ("4", "EmailSender", "Gửi OTP qua Gmail SMTP.", "Nếu thiếu/lỗi SMTP: Development hiển thị OTP thử; môi trường khác báo liên hệ/cấu hình."),
            ("5", "Khách", "Nhập OTP hoặc bấm gửi lại.", "VerifyOtp kiểm tra session, thời hạn, mã; ResendOtp thay mã và gia hạn thêm 10 phút."),
            ("6", "Hệ thống", "Kiểm tra trùng lần cuối và tạo dữ liệu.", "Trong Serializable transaction, lấy số ST lớn nhất +1, insert Student rồi UserAccount Student; commit."),
            ("7", "Hệ thống", "Kết thúc đăng ký.", "Xóa PendingRegistration, báo xác minh thành công và chuyển Login."),
        ],
        "interaction": "Chỉ có tương tác Khách - hệ thống email. Demo dùng một trình duyệt; chuyển sang hộp thư hoặc đọc OTP thử trong Development rồi quay lại trang VerifyOtp.",
        "constraints": [
            ("R1", "Thiếu trường", "Báo 'Vui lòng nhập đầy đủ thông tin' và giữ dữ liệu không nhạy cảm."),
            ("R2", "Username/email/phone trùng", "Không tạo dữ liệu; báo đúng trường/nhóm dữ liệu đã tồn tại."),
            ("R3", "OTP sai", "Giữ trang VerifyOtp và cho nhập lại."),
            ("R4", "OTP/session hết hạn", "Xóa đăng ký tạm và yêu cầu đăng ký lại."),
            ("R5", "Tranh chấp tạo mã ST", "Serializable transaction giảm nguy cơ hai người lấy cùng mã; unique index bảo vệ Code/UserName/Email."),
        ],
        "alternatives": [("Gửi lại OTP", "Mã cũ mất hiệu lực, mã mới có hạn 10 phút."), ("Trùng phát sinh trong lúc chờ OTP", "VerifyOtp phát hiện lại, xóa session và đưa về Register."), ("Thiếu vai trò Student", "Ném InvalidOperationException; đây là lỗi cấu hình hệ thống.")],
        "result": "Ghi Students và UserAccounts; tài khoản IsActive=true. Lưu ý bảo vệ: mật khẩu hiện đang lưu/so sánh dạng rõ, chưa băm và chưa có rate limit.",
        "source": "AuthController.Register, VerifyOtp, ResendOtp, CreateStudentAccountAsync; EmailSender; Program.Session",
    },
    {
        "id": "F05",
        "title": "Đăng nhập, phân quyền, đăng xuất và từ chối truy cập",
        "actors": "Mọi tài khoản; hệ thống cookie authentication.",
        "initiator": "Người dùng mở Login hoặc bấm Đăng xuất.",
        "entry": "Auth/Login; form Logout; Auth/AccessDenied.",
        "prereq": "Đăng nhập cần UserAccount tồn tại và IsActive=true. Đăng xuất cần đã xác thực.",
        "success": "Tạo/xóa cookie EnglishCenter.Auth; gắn claim vai trò và hồ sơ; điều hướng đúng khu vực.",
        "limits": "Cookie 8 giờ, sliding expiration, không persistent; không có khóa tài khoản theo số lần sai.",
        "steps": [
            ("1", "Người dùng", "Nhập username và password.", "AuthController.Login chuẩn hóa username chữ thường và kiểm tra hai trường không rỗng."),
            ("2", "Hệ thống", "Tìm tài khoản đang hoạt động và so sánh mật khẩu.", "Nạp Role; tài khoản không tồn tại, bị khóa hoặc sai mật khẩu đều trả cùng một thông báo."),
            ("3", "Hệ thống", "Tạo ClaimsPrincipal.", "Gắn NameIdentifier, Name, Role, UserName; thêm StudentId hoặc TeacherId nếu có."),
            ("4", "Hệ thống", "Phát cookie đăng nhập.", "IsPersistent=false, AllowRefresh=true; Admin/Staff sang Dashboard, Teacher/Student sang Home/Index."),
            ("5", "Người dùng", "Mở chức năng theo menu.", "[Authorize]/[Authorize(Roles=...)] kiểm tra vai trò; thiếu quyền chuyển AccessDenied/403."),
            ("6", "Người dùng", "Gửi form Đăng xuất.", "AuthController.Logout xóa cookie và quay lại trang chủ."),
        ],
        "interaction": "Không có tác nhân thứ hai. Demo nên mở bốn phiên riêng để quan sát menu và vùng dữ liệu khác nhau cho Admin, Staff, Teacher, Student.",
        "constraints": [("R1", "Thiếu username/password", "Báo yêu cầu nhập đủ, giữ username và returnUrl."), ("R2", "Sai hoặc IsActive=false", "Báo 'Tên đăng nhập hoặc mật khẩu không đúng'."), ("R3", "Đang đăng nhập mà mở Login/Register", "Tự chuyển theo vai trò."), ("R4", "Logout", "Chỉ POST và yêu cầu [Authorize], có antiforgery toàn cục.")],
        "alternatives": [("Admin/Staff", "Đi thẳng Dashboard."), ("Teacher/Student", "Đi trang chủ công khai với menu riêng."), ("Không đủ quyền", "Hiện AccessDenied; không chạy action nghiệp vụ.")],
        "result": "Không đổi bảng nghiệp vụ; đổi trạng thái xác thực ở cookie. Lưu ý bảo vệ: code đang so sánh password dạng rõ.",
        "source": "AuthController.Login, Logout, AccessDenied; Program.Authentication/Authorization",
    },
    {
        "id": "F06",
        "title": "Thông báo trong hệ thống và deep-link",
        "actors": "Người dùng đã đăng nhập; Student; Admin/Staff; NotificationService.",
        "initiator": "Enrollment/Payment tạo thông báo; người nhận bấm biểu tượng chuông hoặc nút Mở.",
        "entry": "Notifications/Index; badge chuông trong _Layout/_AdminLayout.",
        "prereq": "Có claim NameIdentifier; thông báo thuộc đúng UserAccount.",
        "success": "Xem tối đa 50 thông báo mới nhất, đánh dấu đọc và điều hướng tới màn hình xử lý.",
        "limits": "Không có email/push; chỉ lưu trong bảng Notifications; URL chỉ được mở nếu là local URL.",
        "steps": [
            ("1", "Hệ thống", "Một nghiệp vụ đăng ký/thanh toán hoàn tất bước cần thông báo.", "NotificationService tạo Notification cho một user hoặc mọi tài khoản active thuộc role được chỉ định."),
            ("2", "Người nhận", "Quan sát badge chuông và mở danh sách.", "Layout đếm IsRead=false; Index lọc theo UserAccountId, sắp CreatedAt giảm dần, lấy 50."),
            ("3", "Người nhận", "Bấm Mở.", "MarkRead kiểm tra sở hữu, đặt IsRead=true và SaveChanges."),
            ("4", "Hệ thống", "Điều hướng deep-link.", "Thông báo tiêu đề chứa 'đã được duyệt' đi MyPayments; các thông báo khác dùng Url nếu Url.IsLocalUrl."),
            ("5", "Người nhận", "Bấm Đánh dấu đã đọc.", "MarkAllRead ExecuteUpdate toàn bộ thông báo chưa đọc của tài khoản."),
        ],
        "interaction": "Điểm chuyển trình duyệt: sau Student đăng ký/đóng tiền, chuyển sang Admin/Staff để thấy badge; sau duyệt/từ chối, chuyển lại Student để mở thông báo.",
        "constraints": [("R1", "Không có NameIdentifier", "Forbid."), ("R2", "Mở thông báo không thuộc tài khoản", "404 NotFound."), ("R3", "URL ngoài hệ thống", "Không redirect; quay lại Index."), ("R4", "Tài khoản inactive", "Không nhận thông báo mới từ NotifyRolesAsync.")],
        "alternatives": [("Không có thông báo", "Hiện empty state."), ("Đánh dấu tất cả", "Không cần mở từng thông báo; badge về 0 sau khi tải lại.")],
        "result": "Cập nhật Notification.IsRead; deep-link nối trực tiếp tới Enrollments, Payments, MyPayments hoặc MyEnrollments.",
        "source": "NotificationsController; NotificationService; layouts; Courses/Enrollments/Payments controllers",
    },
    {
        "id": "F07",
        "title": "Lưu/bỏ lưu khóa học và quản trị danh sách đã lưu",
        "actors": "Student; Admin; Staff.",
        "initiator": "Student bấm Lưu khóa học; Admin/Staff có thể CRUD bản ghi SavedCourse.",
        "entry": "Courses/Details -> Courses/Save; SavedCourses/MySavedCourses; SavedCourses/Index.",
        "prereq": "Student có claim StudentId; Course tồn tại. Khu quản trị yêu cầu Admin/Staff.",
        "success": "Tạo hoặc xóa SavedCourse đúng cặp Student-Course; trang cá nhân hiển thị khóa mới lưu nhất trước.",
        "limits": "Mỗi Student chỉ lưu một Course một lần nhờ unique index; chức năng không gửi thông báo.",
        "steps": [
            ("1", "Student", "Từ chi tiết khóa học bấm Lưu khóa học.", "CoursesController.Save kiểm tra StudentId và Course, dò cặp đã tồn tại."),
            ("2", "Hệ thống", "Tạo SavedCourse nếu chưa có.", "Ghi StudentId, CourseId, SavedAt; nếu đã có thì không tạo trùng nhưng vẫn báo đã lưu."),
            ("3", "Student", "Mở Khóa học đã lưu.", "MySavedCourses chỉ lấy bản ghi của claim StudentId, Include Course, sắp SavedAt giảm dần."),
            ("4", "Student", "Bấm Bỏ lưu.", "Remove chỉ tìm id thuộc Student đang đăng nhập rồi xóa."),
            ("5", "Admin/Staff", "Mở SavedCourses/Index và Create/Edit/Delete khi cần.", "Controller kiểm tra Student/Course hợp lệ, cặp không trùng; SaveChanges hoặc 404 khi id sai."),
        ],
        "interaction": "Không cần tác nhân khác phản hồi. Khu CRUD Admin/Staff không nằm trong menu chính nhưng action tồn tại và có thể truy cập bằng URL.",
        "constraints": [("R1", "StudentId/Course không hợp lệ", "Save trả 404; form quản trị thêm ModelState error."), ("R2", "Cặp đã tồn tại", "Student Save là idempotent; CRUD quản trị báo đã lưu khóa này."), ("R3", "Remove bản ghi người khác", "404 NotFound."), ("R4", "Xóa Student hoặc Course", "SavedCourse bị cascade theo cấu hình quan hệ.")],
        "alternatives": [("Khách", "Challenge tới Login."), ("Đăng nhập sai vai trò", "Forbid qua kiểm tra thủ công hoặc [Authorize].")],
        "result": "Ghi/xóa SavedCourses; không ảnh hưởng Enrollment hoặc Payment.",
        "source": "CoursesController.Save; SavedCoursesController; DbContext unique index SavedCourse",
    },
    {
        "id": "F08",
        "title": "Đăng ký lớp, duyệt/xếp lớp và hủy đăng ký",
        "actors": "Student khởi tạo; Admin/Staff xử lý; hệ thống thông báo; Student nhận kết quả.",
        "initiator": "Student bấm Đăng ký lớp này ở Courses/Details.",
        "entry": "Courses/Register -> Enrollments/MyEnrollments -> Enrollments/Index/UpdateStatus.",
        "prereq": "Student đăng nhập; Course/Class tồn tại; lớp còn nhận đăng ký và chưa đủ sĩ số; chưa có Enrollment khác chưa hủy cho cùng Course.",
        "success": "Enrollment=Approved gắn CourseClass; tạo Payment=Unpaid; Student được thông báo tới MyPayments.",
        "limits": "Sĩ số chỉ đếm Approved; thao tác đăng ký và duyệt dùng Serializable transaction; một đăng ký active/Student/Course.",
        "steps": [
            ("1", "Student", "Chọn một lớp còn chỗ và bấm Đăng ký lớp này.", "CoursesController.Register lấy StudentId, bắt đầu Serializable transaction."),
            ("2", "Hệ thống", "Kiểm tra Course, đăng ký trùng, lớp thuộc Course, CanRegister và Capacity.", "Nếu hợp lệ tạo Enrollment=Pending, lưu CourseClassId đã chọn và RegisteredAt."),
            ("3", "Hệ thống", "Thông báo Admin và Staff.", "NotifyRolesAsync gửi 'Đăng ký khóa học mới' kèm link Enrollments/Index; commit."),
            ("4", "Admin/Staff", "Mở thông báo hoặc Duyệt đăng ký, tìm/lọc hồ sơ.", "Enrollments.Index nạp Student/Course/Class, trạng thái, lớp và sĩ số; 10 bản ghi/trang."),
            ("5", "Admin/Staff", "Chọn lớp của đúng khóa và bấm Duyệt.", "UpdateStatus mở Serializable transaction; kiểm tra lớp còn mở và còn chỗ, loại chính enrollment khi đếm."),
            ("6", "Hệ thống", "Cập nhật Enrollment và học phí.", "Đặt Approved + CourseClassId; nếu chưa có Payment thì tạo theo Course.Tuition, PaidAmount=0, Status=Unpaid, method=Cash."),
            ("7", "Hệ thống", "Thông báo Student.", "Tạo notification 'đã được duyệt' với link MyPayments; commit."),
            ("8", "Student", "Mở thông báo và xem Khóa học của tôi/Học phí.", "MyEnrollments chỉ lấy dữ liệu của chính Student; hiển thị lớp, Teacher, trạng thái và Payment."),
        ],
        "interaction": "Demo 2 trình duyệt: A-Student đăng ký -> B-Admin/Staff thấy badge và duyệt -> A-Student thấy badge, mở thông báo và chuyển thẳng tới Học phí.",
        "constraints": [
            ("R1", "Đã có Enrollment khác Status != Cancelled", "Chặn và báo đã đăng ký khóa này; unique filtered index là lớp bảo vệ cuối."),
            ("R2", "Lớp không thuộc khóa/đã khóa/đóng/đầy", "Không tạo hoặc không duyệt; trả thông báo cụ thể."),
            ("R3", "Duyệt không chọn lớp", "Báo yêu cầu chọn lớp."),
            ("R4", "Student tự hủy", "Chỉ cho Enrollment của mình đang Pending; chuyển Cancelled và bỏ CourseClassId."),
            ("R5", "Admin/Staff hủy", "Nếu Payment.PaidAmount>0 hoặc Paid thì chặn; nếu chưa thu thì xóa Payment và hủy Enrollment."),
            ("R6", "Tranh chấp chỗ cuối", "Serializable transaction bao trọn kiểm tra sĩ số và ghi; DB unique index ngăn đăng ký active trùng."),
        ],
        "alternatives": [("Admin/Staff bấm Hủy", "Enrollment=Cancelled, CourseClassId=null; Student nhận thông báo tới MyEnrollments."), ("Student hủy Pending", "Không cần chờ Admin; có thể đăng ký lại cùng Course vì index bỏ qua Cancelled."), ("Đã phát sinh tiền", "Không hủy trực tiếp; phải xử lý học phí/hoàn tiền ngoài luồng hiện tại.")],
        "result": "Ghi Enrollments, Notifications và có thể Payments; ảnh hưởng sĩ số lớp, lịch học, học phí, bài giảng, điểm danh và điểm số về sau.",
        "source": "CoursesController.Register; EnrollmentsController.Index, UpdateStatus, MyEnrollments, Cancel; DbContext indexes",
    },
    {
        "id": "F09",
        "title": "Thanh toán học phí, QR, duyệt/từ chối và hóa đơn",
        "actors": "Student khởi tạo thanh toán; Admin/Staff duyệt hoặc ghi nhận; NotificationService; dịch vụ ảnh VietQR.",
        "initiator": "Student mở Học phí và chọn số tiền/phương thức; hoặc Admin/Staff ghi nhận trực tiếp.",
        "entry": "Payments/MyPayments -> Payments/QR hoặc Payments/Pay -> Payments/Index -> Reports/InvoicePdf.",
        "prereq": "Enrollment=Approved; Payment chưa Cancelled/Paid; còn công nợ; người thao tác đúng quyền/sở hữu.",
        "success": "Transaction=Approved; PaidAmount tăng; Payment=PartiallyPaid hoặc Paid; Student nhận thông báo; Paid được xem/xuất hóa đơn PDF.",
        "limits": "Cho đóng một phần; tổng giao dịch Pending không vượt số còn lại; QR dùng tài khoản TCB cố định và không có webhook xác nhận ngân hàng.",
        "steps": [
            ("1", "Student", "Mở MyPayments, nhập số tiền 1..còn lại và chọn phương thức.", "Trang hiển thị Amount, PaidAmount, remaining và lịch sử transaction."),
            ("2", "Student", "Nếu BankTransfer, trình duyệt chuyển tới trang QR.", "QR action kiểm tra sở hữu, chọn amount hợp lệ hoặc remaining; tạo URL ảnh VietQR với nội dung HP_{PaymentId}."),
            ("3", "Student", "Quét QR rồi bấm Tôi đã thanh toán; Card/Cash bấm Thanh toán trực tiếp.", "Payments.Pay chạy Serializable transaction; kiểm tra trạng thái, amount và tổng Pending."),
            ("4", "Hệ thống", "Tạo yêu cầu chờ duyệt.", "Tạo PaymentTransaction=Pending, RecordedBy=Student; đặt Payment.Status=PendingApproval; thông báo mọi Admin/Staff kèm link Payments/Index."),
            ("5", "Admin/Staff", "Mở Payments/Index, xem giao dịch Pending và bấm Duyệt hoặc Từ chối.", "Danh sách ưu tiên khoản có Pending; hiển thị từng transaction chờ xử lý."),
            ("6", "Hệ thống", "Khi duyệt, kiểm tra transaction còn Pending và số tiền không vượt công nợ.", "Trong Serializable transaction, cộng PaidAmount, cập nhật method/date/status; transaction=Approved, ApprovedAt/By."),
            ("7", "Hệ thống", "Thông báo Student.", "Nếu đã Paid, deep-link tới InvoicePdf; nếu mới PartiallyPaid, link MyPayments."),
            ("8", "Student", "Mở thông báo hoặc bấm Xem hóa đơn.", "Reports.InvoicePdf kiểm tra đúng chủ sở hữu và Payment=Paid rồi sinh PDF."),
            ("9", "Admin/Staff", "Lối tắt: nhập tổng PaidAmount và phương thức, bấm Lưu.", "RecordPayment điều chỉnh trực tiếp tổng đã đóng; tạo transaction Approved bằng trị tuyệt đối phần chênh lệch và cập nhật Payment.Status."),
        ],
        "interaction": "Demo 2 trình duyệt: A-Student tạo thanh toán -> B-Admin/Staff thấy badge và Pending -> B duyệt/từ chối -> A nhận badge, kiểm tra trạng thái và hóa đơn.",
        "constraints": [
            ("R1", "amount <=0 hoặc > remaining", "Chặn và báo miền số tiền hợp lệ."),
            ("R2", "amount + tổng Pending > remaining", "Chặn để tránh nhiều yêu cầu vượt công nợ."),
            ("R3", "Transaction đã xử lý", "Không duyệt/từ chối lần hai."),
            ("R4", "Duyệt vượt công nợ hiện tại", "Chặn; transaction vẫn Pending để xử lý lại/từ chối."),
            ("R5", "Từ chối", "Transaction=Rejected; Payment giữ PendingApproval nếu còn transaction Pending khác, nếu không tính lại theo PaidAmount."),
            ("R6", "Xuất hóa đơn", "Chỉ Payment=Paid; Student chỉ xem hóa đơn của mình; Admin/Staff được tải file có tên."),
        ],
        "alternatives": [("Đóng một phần", "Payment=PartiallyPaid sau khi duyệt; Student tiếp tục đóng phần còn lại."), ("Admin/Staff ghi tiền mặt", "Không cần chờ Student; transaction được tạo Approved ngay."), ("Từ chối", "PaidAmount không tăng; Student nhận thông báo kiểm tra lại."), ("QR không tải", "Do dịch vụ ảnh ngoài; có thể dùng phương thức khác hoặc thông tin tài khoản hiển thị trên trang.")],
        "result": "Ghi Payments, PaymentTransactions, Notifications; giao dịch Approved ảnh hưởng Dashboard, báo cáo doanh thu quý và quyền xuất hóa đơn.",
        "source": "PaymentsController; ReportsController.InvoicePdf; Views/Payments/MyPayments.cshtml, QR.cshtml, Index.cshtml",
    },
    {
        "id": "F10",
        "title": "Cập nhật hồ sơ cá nhân học viên và giáo viên",
        "actors": "Student hoặc Teacher; UserAccount liên kết.",
        "initiator": "Người dùng mở menu tài khoản và chọn Hồ sơ cá nhân/giảng viên.",
        "entry": "Students/Profile hoặc Teachers/Profile.",
        "prereq": "Đăng nhập đúng role và có claim StudentId/TeacherId; hồ sơ tương ứng tồn tại.",
        "success": "Cập nhật hồ sơ chuyên môn/cá nhân và đồng bộ FullName, Email, Phone sang UserAccount liên kết.",
        "limits": "Không đổi Code, username, role hoặc password; không có upload ảnh hồ sơ.",
        "steps": [
            ("1", "Student/Teacher", "Mở trang hồ sơ.", "Action kiểm tra role + claim thủ công; chỉ đọc đúng Id trong claim."),
            ("2", "Student", "Sửa họ tên, email, điện thoại, ngày sinh, địa chỉ và gửi.", "Gắn lại Code cũ, TryValidateModel, kiểm tra trùng Student/UserAccount; cập nhật Student."),
            ("3", "Teacher", "Sửa họ tên, email, điện thoại, chuyên môn, trình độ, chứng chỉ và gửi.", "Gắn lại Code cũ, validate độ dài/định dạng/trùng; cập nhật Teacher."),
            ("4", "Hệ thống", "Đồng bộ tài khoản.", "Nếu UserAccount liên kết tồn tại, cập nhật FullName, Email, Phone trong cùng SaveChanges."),
            ("5", "Hệ thống", "Thông báo thành công.", "Redirect lại Profile để hiển thị dữ liệu mới."),
        ],
        "interaction": "Không cần Admin/Staff duyệt. Thay đổi tên/email/phone được phản ánh ở tài khoản nhưng claim Name trong cookie chỉ cập nhật sau lần đăng nhập tiếp theo.",
        "constraints": [("R1", "Student", "Code STxx, tên <=100, email <=150 hợp lệ, phone <=20 hợp lệ, tuổi 5-100, địa chỉ <=250."), ("R2", "Teacher", "Code TCxx, tên <=100, email/phone bắt buộc, chuyên môn <=150, trình độ <=200, chứng chỉ <=500."), ("R3", "Email/phone trùng", "Giữ form và hiển thị lỗi, không SaveChanges."), ("R4", "Sai role/claim", "Challenge nếu chưa đăng nhập, Forbid nếu đã đăng nhập sai vai trò.")],
        "alternatives": [("Hồ sơ không tồn tại", "404 NotFound."), ("Không có UserAccount liên kết", "Vẫn cập nhật hồ sơ, bỏ qua bước đồng bộ.")],
        "result": "Ghi Students hoặc Teachers và có thể UserAccounts; ảnh hưởng tên hiển thị ở nhiều danh sách sau khi đăng nhập lại.",
        "source": "StudentsController.Profile; TeachersController.Profile; Student/Teacher models",
    },
    {
        "id": "F11",
        "title": "Bảng điều khiển vận hành",
        "actors": "Admin, Staff.",
        "initiator": "Đăng nhập Admin/Staff hoặc bấm Bảng điều khiển.",
        "entry": "Home/Dashboard.",
        "prereq": "Đã đăng nhập role Admin hoặc Staff.",
        "success": "Hiển thị chỉ số quy mô, đăng ký, học phí, tiến độ thu và 8 đăng ký gần đây đã duyệt + hoàn tất học phí.",
        "limits": "Số liệu tính trực tiếp từ database tại thời điểm tải trang; không lưu snapshot; không có chọn khoảng thời gian.",
        "steps": [
            ("1", "Admin/Staff", "Mở Dashboard.", "Đếm Students, Teachers, Courses, CourseClasses, Pending/Approved Enrollments."),
            ("2", "Hệ thống", "Tính tài chính.", "Cộng Payment.Amount và PaidAmount với Enrollment không Cancelled; view tính tỷ lệ thu tối đa 100%."),
            ("3", "Hệ thống", "Nạp hoạt động gần đây.", "Lấy 8 Enrollment Approved có Payment Paid hoặc PaidAmount>=Amount, sắp RegisteredAt giảm dần."),
            ("4", "Admin/Staff", "Dùng thao tác nhanh.", "Chuyển tới duyệt đăng ký, tạo lớp hoặc báo cáo doanh thu."),
        ],
        "interaction": "Không thông báo cho tác nhân khác; các chỉ số thay đổi sau khi quy trình đăng ký/thanh toán được xử lý và trang được tải lại.",
        "constraints": [("R1", "Không đúng role", "[Authorize(Roles='Admin,Staff')] chặn."), ("R2", "Không có dữ liệu", "Các tổng bằng 0; danh sách gần đây rỗng."), ("R3", "Enrollment Cancelled", "Không tính vào Expected/CollectedTuition.")],
        "alternatives": [("Bấm hồ sơ chờ duyệt", "Mở Enrollments/Index."), ("Bấm xuất báo cáo", "Mở F22 báo cáo doanh thu quý.")],
        "result": "Chỉ đọc dữ liệu; là màn hình điều phối các chức năng vận hành.",
        "source": "HomeController.Dashboard; Views/Home/Dashboard.cshtml",
    },
    {
        "id": "F12",
        "title": "Quản lý khóa học (CRUD)",
        "actors": "Admin, Staff.",
        "initiator": "Admin/Staff mở Khóa học trong khu vận hành.",
        "entry": "Courses/Index -> Create/Edit/Details/Delete.",
        "prereq": "Đăng nhập Admin/Staff.",
        "success": "Thêm/sửa/xóa Course hợp lệ; danh sách và trang công khai phản ánh thay đổi.",
        "limits": "6 bản ghi/trang; ImageUrl là chuỗi đường dẫn, không upload ảnh; xóa bị chặn khi có dữ liệu phụ thuộc.",
        "steps": [
            ("1", "Admin/Staff", "Tìm theo mã/tên, trình độ, ngày khai giảng.", "Dùng chung Courses.Index với F02 nhưng hiển thị bảng quản trị và nút Thêm/Sửa/Xóa."),
            ("2", "Admin/Staff", "Bấm Thêm mới, nhập dữ liệu và gửi.", "Create chuẩn hóa Code viết hoa, trim trường, TryValidateModel, kiểm tra Code trùng, insert Course."),
            ("3", "Admin/Staff", "Bấm Sửa và gửi dữ liệu mới.", "Edit kiểm tra id khớp, chuẩn hóa/validate và kiểm tra Code trùng trừ bản ghi hiện tại."),
            ("4", "Admin/Staff", "Bấm Xóa, xem xác nhận và gửi POST.", "DeleteConfirmed thử Remove; DbUpdateException được đổi thành thông báo không thể xóa do liên quan."),
        ],
        "interaction": "Không có duyệt hai bước; thay đổi xuất hiện ngay cho khách/Student. Không gửi notification tới người học.",
        "constraints": [("R1", "Code", "Bắt buộc, <=20, regex ^CR\\d{2,}$, unique."), ("R2", "Name/Level/Duration/Description", "Bắt buộc; giới hạn lần lượt 150/50/80/2000 ký tự."), ("R3", "Tuition", "0..1.000.000.000, precision 18,2."), ("R4", "ImageUrl", "Tối đa 500 ký tự."), ("R5", "Xóa", "Bị Restrict nếu có CourseClass, Enrollment hoặc CourseLecture; SavedCourse cascade nhưng không vượt được các quan hệ Restrict khác.")],
        "alternatives": [("ModelState lỗi", "Giữ model và hiển thị validation; không ghi."), ("id không tồn tại/không khớp", "404 NotFound."), ("Xung đột FK khi xóa", "Giữ Course và hiện thông báo lỗi.")],
        "result": "Ghi Courses; ảnh hưởng danh mục công khai, lớp, đăng ký, học phí và bài giảng.",
        "source": "CoursesController.Create/Edit/Delete; Course model; DbContext relationships",
    },
    {
        "id": "F13",
        "title": "Quản lý lớp học và lịch khai giảng (CRUD)",
        "actors": "Admin, Staff.",
        "initiator": "Admin/Staff mở Lớp học & lịch khai giảng.",
        "entry": "CourseClasses/Index -> Create/Edit/Details/Delete.",
        "prereq": "Có Course và Teacher hợp lệ; đăng nhập Admin/Staff.",
        "success": "CourseClass được tạo/cập nhật/xóa hợp lệ và xuất hiện trong lịch công khai.",
        "limits": "10 lớp/trang; chưa kiểm tra trùng lịch phòng hoặc lịch giáo viên; trạng thái hiệu lực có phần được suy ra theo ngày.",
        "steps": [
            ("1", "Admin/Staff", "Tìm theo lớp/khóa/giáo viên hoặc StartDate.", "Index Include Course, Teacher, Enrollment; phân trang 10."),
            ("2", "Admin/Staff", "Bấm Thêm lớp, chọn Course/Teacher, nhập Code, Room, Schedule, ngày, Status, Capacity.", "Create chuẩn hóa, validate Code/FK/ngày/lịch rồi insert."),
            ("3", "Admin/Staff", "Mở Chi tiết.", "Nạp Course, Teacher và toàn bộ Enrollment kèm Student để xem lớp."),
            ("4", "Admin/Staff", "Sửa lớp.", "Edit kiểm tra id, chạy lại toàn bộ validation và update."),
            ("5", "Admin/Staff", "Xóa lớp.", "Remove; nếu có Enrollment, Score hoặc Attendance thì catch DbUpdateException và giữ dữ liệu."),
        ],
        "interaction": "Thay đổi lớp ảnh hưởng ngay đến đăng ký, lịch Student/Teacher và điểm danh; code không gửi thông báo tự động khi đổi lịch.",
        "constraints": [("R1", "Code", "Bắt buộc, <=20, regex ^CL\\d{2,}$, unique."), ("R2", "Course/Teacher", "Phải tồn tại."), ("R3", "Room/Schedule", "Bắt buộc, <=80/150; Schedule phải phân tích được ít nhất một thứ học."), ("R4", "Ngày", "EndDate nếu có phải >= StartDate."), ("R5", "Capacity", "1..500."), ("R6", "CanRegister", "Locked/Closed không nhận đăng ký; hết EndDate thành Closed; StartDate tương lai thành Upcoming.")],
        "alternatives": [("Lịch không có Thứ 2..7/CN", "ModelState lỗi với ví dụ định dạng."), ("Lớp có dữ liệu liên quan", "Không xóa."), ("Trùng giờ/phòng/giáo viên", "Hiện code chưa chặn; cần nêu khi bảo vệ.")],
        "result": "Ghi CourseClasses; ảnh hưởng Enrollment, Schedule, Roster, Attendance, Score.",
        "source": "CourseClassesController; CourseClass model; DbContext",
    },
    {
        "id": "F14",
        "title": "Quản lý hồ sơ học viên (CRUD)",
        "actors": "Admin, Staff.",
        "initiator": "Admin/Staff mở Học viên.",
        "entry": "Students/Index -> Create/Edit/Details/Delete.",
        "prereq": "Đăng nhập Admin/Staff.",
        "success": "Thêm/sửa/xóa Student hợp lệ; thông tin được dùng trong tài khoản, đăng ký và học tập.",
        "limits": "10 học viên/trang; tạo Student không tự tạo UserAccount; xóa bị chặn bởi dữ liệu học tập/tài chính.",
        "steps": [
            ("1", "Admin/Staff", "Tìm theo Code, tên, email, phone.", "Students.Index lọc, sắp Code và phân trang."),
            ("2", "Admin/Staff", "Thêm học viên.", "Create trim/viết hoa Code, TryValidateModel, kiểm tra Code/email/phone trùng rồi insert."),
            ("3", "Admin/Staff", "Sửa học viên.", "Edit kiểm tra id, validate lại và update toàn entity Student."),
            ("4", "Admin/Staff", "Xóa học viên.", "UserAccount liên kết được SetNull; SavedCourses cascade; các quan hệ Restrict có thể chặn và controller báo lỗi."),
        ],
        "interaction": "Không có duyệt hoặc notification. Nếu muốn học viên đăng nhập, Admin phải tạo/liên kết UserAccount ở F16 hoặc học viên tự đăng ký ở F04.",
        "constraints": [("R1", "Code", "ST + ít nhất 2 chữ số; <=20; unique."), ("R2", "Tên/email/phone/address", "Bắt buộc; email/phone đúng định dạng; độ dài 100/150/20/250."), ("R3", "Tuổi", "5..100 tại ngày hiện tại."), ("R4", "Trùng", "Kiểm tra trong Students và UserAccounts để tránh email/phone thuộc tài khoản khác."), ("R5", "Xóa", "Chặn khi có Enrollment, Payment, PaymentTransaction, Score hoặc AttendanceRecord.")],
        "alternatives": [("Dữ liệu lỗi", "Giữ form và hiển thị ModelState."), ("id sai", "404 NotFound."), ("Có dữ liệu liên quan", "Giữ Student, hiện thông báo không thể xóa.")],
        "result": "Ghi Students; có thể làm UserAccount.StudentId thành null khi xóa thành công.",
        "source": "StudentsController CRUD; Student model; DbContext delete behaviors",
    },
    {
        "id": "F15",
        "title": "Quản lý hồ sơ giáo viên (CRUD)",
        "actors": "Admin, Staff.",
        "initiator": "Admin/Staff mở Giáo viên.",
        "entry": "Teachers/Index -> Create/Edit/Details/Delete.",
        "prereq": "Đăng nhập Admin/Staff.",
        "success": "Thêm/sửa/xóa Teacher hợp lệ; hồ sơ dùng cho lớp và bài giảng.",
        "limits": "10 giáo viên/trang; tạo Teacher không tự tạo UserAccount; xóa bị chặn khi có lớp/bài giảng.",
        "steps": [
            ("1", "Admin/Staff", "Tìm theo Code, tên, email, phone, chuyên môn, chứng chỉ.", "Teachers.Index lọc và phân trang."),
            ("2", "Admin/Staff", "Thêm giáo viên.", "Create chuẩn hóa Code, trim trường, validate và kiểm tra trùng rồi insert."),
            ("3", "Admin/Staff", "Sửa giáo viên.", "Edit kiểm tra id và update sau validation."),
            ("4", "Admin/Staff", "Xóa giáo viên.", "UserAccount liên kết SetNull; CourseClass/CourseLecture Restrict có thể chặn, controller báo lỗi."),
        ],
        "interaction": "Không thông báo tự động. Tài khoản Teacher được tạo/liên kết riêng ở F16.",
        "constraints": [("R1", "Code", "TC + ít nhất 2 chữ số; <=20; unique."), ("R2", "Tên/email/phone/chuyên môn", "Bắt buộc; giới hạn 100/150/20/150; email/phone đúng định dạng."), ("R3", "Trình độ/chứng chỉ", "Không bắt buộc; tối đa 200/500."), ("R4", "Trùng email/phone", "Kiểm tra Teachers và UserAccounts."), ("R5", "Xóa", "Không được nếu có CourseClass hoặc CourseLecture.")],
        "alternatives": [("ModelState lỗi", "Giữ form."), ("id sai", "404."), ("Có lớp/bài giảng", "Không xóa và hiện lỗi.")],
        "result": "Ghi Teachers; ảnh hưởng lớp, lịch khai giảng, roster và bài giảng.",
        "source": "TeachersController CRUD; Teacher model; DbContext",
    },
    {
        "id": "F16",
        "title": "Quản lý tài khoản và kích hoạt",
        "actors": "Admin.",
        "initiator": "Admin mở Tài khoản.",
        "entry": "UserAccounts/Index -> Create/Edit/Details/Delete.",
        "prereq": "Đăng nhập Admin; Roles tồn tại; hồ sơ Student/Teacher tồn tại nếu liên kết.",
        "success": "UserAccount hợp lệ, đúng Role và liên kết 1-1; có thể bật/tắt IsActive hoặc đổi mật khẩu.",
        "limits": "10 tài khoản/trang; chỉ Admin; không được xóa chính tài khoản đang đăng nhập.",
        "steps": [
            ("1", "Admin", "Tìm theo username, tên, email, phone hoặc vai trò.", "Index Include Role, Student, Teacher; phân trang."),
            ("2", "Admin", "Bấm Tạo tài khoản, nhập thông tin, chọn role và hồ sơ liên kết.", "Create chuẩn hóa, yêu cầu password, kiểm tra trùng và quy tắc liên kết; insert với CreatedAt."),
            ("3", "Admin", "Sửa tài khoản.", "Mật khẩu để trống thì giữ mật khẩu cũ; cập nhật IsActive, role và liên kết sau validation."),
            ("4", "Admin", "Tắt IsActive.", "Tài khoản vẫn tồn tại nhưng Login không tìm thấy; cookie đã phát trước đó không bị thu hồi ngay trong code."),
            ("5", "Admin", "Xóa tài khoản khác.", "Kiểm tra NameIdentifier không trùng id hiện tại; remove, Notification cascade."),
        ],
        "interaction": "Không gửi thông báo. Người bị tắt chỉ thấy không đăng nhập được ở lần đăng nhập sau; hệ thống chưa có cơ chế ép đăng xuất tức thời.",
        "constraints": [("R1", "Username", "Bắt buộc, 3..50, chuẩn hóa chữ thường, unique."), ("R2", "Password", "Create bắt buộc; Edit tùy chọn. Model không đặt min length cho Admin-created account."), ("R3", "Email/phone", "Đúng định dạng và không trùng UserAccount khác."), ("R4", "Role Student", "Phải có StudentId; TeacherId bị xóa; một Student chỉ có một account."), ("R5", "Role Teacher", "Phải có TeacherId; StudentId bị xóa; một Teacher chỉ có một account."), ("R6", "Role khác", "Cả StudentId và TeacherId phải null."), ("R7", "Đồng thời", "Unique indexes và catch DbUpdateException xử lý trường hợp hai phiên cùng tạo tài khoản/liên kết.")],
        "alternatives": [("Role/hồ sơ không hợp lệ", "ModelState error, không lưu."), ("Xóa tài khoản hiện tại", "Bị chặn."), ("Tạo đồng thời bị giành username/hồ sơ", "Catch và báo dữ liệu vừa được sử dụng.")],
        "result": "Ghi UserAccounts; ảnh hưởng đăng nhập, menu, quyền và nhận Notification. Lưu ý: password chưa băm.",
        "source": "UserAccountsController; UserAccount model; DbContext indexes/relations",
    },
    {
        "id": "F17",
        "title": "Quản lý vai trò",
        "actors": "Admin.",
        "initiator": "Admin mở Vai trò.",
        "entry": "Roles/Index -> Create/Edit/Details/Delete.",
        "prereq": "Đăng nhập Admin.",
        "success": "Role được thêm/sửa/xóa nếu hợp lệ và chưa được tài khoản sử dụng.",
        "limits": "Tên role được dùng trực tiếp trong [Authorize] và menu; sửa/xóa role hệ thống có thể làm lệch phân quyền.",
        "steps": [
            ("1", "Admin", "Xem danh sách/chi tiết role.", "RolesController đọc bảng Roles."),
            ("2", "Admin", "Thêm Name và DisplayName.", "Trim, TryValidateModel, kiểm tra Name unique, insert."),
            ("3", "Admin", "Sửa role.", "Kiểm tra id, validate Name unique, update; xử lý concurrency nếu bản ghi đã bị xóa."),
            ("4", "Admin", "Xóa role.", "Remove; quan hệ UserAccount.RoleId Restrict khiến DbUpdateException nếu đang dùng."),
        ],
        "interaction": "Không thông báo. Role Name mới không tự có menu/chính sách tương ứng nếu code [Authorize] chưa hỗ trợ.",
        "constraints": [("R1", "Name", "Bắt buộc, <=30, unique."), ("R2", "DisplayName", "Bắt buộc, <=80."), ("R3", "Đang có account", "Không xóa."), ("R4", "Tên chuẩn Admin/Staff/Teacher/Student", "Được hard-code trong nhiều Controller/View; đổi tên có thể mất quyền.")],
        "alternatives": [("Name trùng", "ModelState error."), ("Concurrency", "404 nếu role đã mất; ngoại lệ khác được ném lại."), ("Xóa role đang dùng", "Giữ role và báo lỗi.")],
        "result": "Ghi Roles; ảnh hưởng trực tiếp phân quyền và hiển thị menu.",
        "source": "RolesController; Role model; DbContext Role.Name unique và DeleteBehavior.Restrict",
    },
    {
        "id": "F18",
        "title": "Lớp giảng dạy, danh sách lớp và lịch học cá nhân",
        "actors": "Teacher; Student.",
        "initiator": "Teacher mở Lớp giảng dạy; Student mở Lịch học.",
        "entry": "CourseClasses/MyClasses, ClassRoster, MySchedule.",
        "prereq": "TeacherId/StudentId hợp lệ; Enrollment phải Approved để xuất hiện trong roster/lịch Student.",
        "success": "Teacher chỉ xem lớp mình dạy và roster đã duyệt; Student chỉ xem lịch lớp đã duyệt của mình.",
        "limits": "Chỉ đọc; MyClasses hiển thị mọi trạng thái lớp của giáo viên; lịch Student sắp theo StartDate rồi Schedule.",
        "steps": [
            ("1", "Teacher", "Mở Lớp giảng dạy.", "MyClasses lọc CourseClass.TeacherId theo claim, Include Course/Enrollments và đếm Approved ở View."),
            ("2", "Teacher", "Bấm Danh sách học viên.", "ClassRoster kiểm tra lớp thuộc Teacher; chỉ Include Enrollment=Approved và Student."),
            ("3", "Teacher", "Từ lớp chọn Điểm danh hoặc Nhập điểm.", "Đi tới F19/F20 với classId."),
            ("4", "Student", "Mở Lịch học.", "MySchedule lọc Enrollment của Student, Approved, có CourseClass; Include Course/Class/Teacher."),
            ("5", "Hệ thống", "Hiển thị lịch.", "Mỗi thẻ có ngày bắt đầu, mã lớp, khóa, lịch, phòng và giáo viên."),
        ],
        "interaction": "Admin/Staff duyệt/xếp lớp ở F08 là điều kiện để Teacher thấy Student trong roster và Student thấy lịch. Khi demo, duyệt ở trình duyệt Admin rồi tải lại hai phía.",
        "constraints": [("R1", "Thiếu claim", "Forbid."), ("R2", "Teacher mở lớp người khác", "404 NotFound."), ("R3", "Enrollment Pending/Cancelled", "Không xuất hiện trong roster hoặc MySchedule."), ("R4", "Student chưa có lịch", "Hiện thông báo chưa có lịch được xếp.")],
        "alternatives": [("Lớp không có học viên Approved", "Roster rỗng."), ("Enrollment Approved nhưng CourseClassId null", "Không hiện trong MySchedule.")],
        "result": "Không ghi dữ liệu; kết nối luồng duyệt đăng ký với điểm danh, điểm và học liệu.",
        "source": "CourseClassesController.MyClasses, ClassRoster, MySchedule; related views",
    },
    {
        "id": "F19",
        "title": "Điểm danh theo lớp và tra cứu chuyên cần",
        "actors": "Teacher nhập; Student xem; Admin/Staff tổng hợp.",
        "initiator": "Teacher từ MyClasses bấm Điểm danh.",
        "entry": "AttendanceRecords/Manage -> SaveAll/Save; MyAttendance; AttendanceRecords/Index.",
        "prereq": "Teacher sở hữu lớp; Student trong Enrollment=Approved; ngày thuộc thời gian và thứ học của lớp.",
        "success": "Mỗi Student/Class/StudyDate có một AttendanceRecord; Student xem được số buổi có mặt/vắng và lịch sử.",
        "limits": "Ghi chú <=500; unique composite; parser lịch chỉ dựa phần trước dấu phẩy và các token Thứ 2..7/CN.",
        "steps": [
            ("1", "Teacher", "Mở Manage cho lớp và chọn ngày.", "Kiểm tra CourseClass.TeacherId; ValidateStudyDate và nạp danh sách Enrollment Approved + bản ghi ngày đó."),
            ("2", "Teacher", "Chọn Có mặt/Vắng, nhập ghi chú cho từng Student, bấm Lưu tất cả.", "SaveAll xác thực lớp/ngày, danh sách Student phải là tập con approved, kiểm tra ghi chú."),
            ("3", "Hệ thống", "Upsert từng AttendanceRecord.", "Tìm composite Student/Class/Date; tạo nếu thiếu hoặc cập nhật IsPresent/Note; SaveChanges một lần."),
            ("4", "Student", "Mở Điểm danh của tôi, chọn lớp và khoảng ngày.", "MyAttendance chỉ cho lớp Approved của Student, lọc from/to, đếm present/absent."),
            ("5", "Admin/Staff", "Mở Lịch sử điểm danh và tìm/lọc.", "Index lọc theo Student/Class/Course keyword và StudyDate; phân trang 10."),
        ],
        "interaction": "Sau Teacher lưu, chuyển sang trình duyệt Student và tải lại MyAttendance để thấy bản ghi; Admin/Staff chỉ quan sát, không có action sửa từ Controller hiện tại.",
        "constraints": [("R1", "Ngày trước StartDate", "Chặn với thông báo ngày bắt đầu."), ("R2", "Ngày sau EndDate", "Chặn với thông báo ngày kết thúc."), ("R3", "Sai thứ trong Schedule", "Chặn và nêu lịch lớp."), ("R4", "Student không Approved hoặc không thuộc lớp", "404."), ("R5", "Ghi chú >500", "Không lưu batch."), ("R6", "Đồng thời", "Unique index ngăn hai bản ghi cùng composite; code chưa bắt DbUpdateException cho va chạm đồng thời.")],
        "alternatives": [("Lưu một Student", "Save action có cùng kiểm tra và upsert."), ("Lớp chưa có học viên", "SaveAll với mảng rỗng trả 404; UI hiển thị bảng rỗng."), ("Student chưa có lớp Approved", "Hiện thông báo chưa có lớp và số liệu 0.")],
        "result": "Ghi AttendanceRecords; ảnh hưởng màn hình chuyên cần Student và bảng tổng hợp Admin/Staff.",
        "source": "AttendanceRecordsController; AttendanceRecord model; DbContext unique index",
    },
    {
        "id": "F20",
        "title": "Nhập điểm, xem kết quả và xuất phiếu kết quả",
        "actors": "Teacher nhập; Student xem; Admin/Staff tổng hợp và xuất PDF.",
        "initiator": "Teacher từ MyClasses bấm Nhập điểm.",
        "entry": "Scores/Manage -> Save; Scores/MyScores; Scores/Index -> Reports/StudentResultPdf.",
        "prereq": "Teacher sở hữu lớp; Student có Enrollment=Approved trong lớp.",
        "success": "Một Score/Student/Class được tạo/cập nhật; tính trung bình và Đạt/Chưa đạt; Admin/Staff xuất PDF.",
        "limits": "Điểm 0..10; nhận xét <=500; trung bình=(giữa kỳ+cuối kỳ)/2, làm tròn 2 chữ số; đạt từ 5.",
        "steps": [
            ("1", "Teacher", "Mở Manage của lớp.", "Kiểm tra lớp thuộc Teacher; nạp Enrollment Approved và Score đã có."),
            ("2", "Teacher", "Nhập giữa kỳ, cuối kỳ, nhận xét cho một Student và bấm Lưu.", "Scores.Save kiểm tra quyền lớp + Enrollment, miền điểm, độ dài nhận xét."),
            ("3", "Hệ thống", "Upsert Score.", "Tìm theo StudentId/CourseClassId; tạo nếu thiếu hoặc cập nhật; SaveChanges."),
            ("4", "Student", "Mở Kết quả học tập.", "MyScores chỉ lấy Score của Student; View tính AverageScore và Result."),
            ("5", "Admin/Staff", "Mở danh sách, tìm theo học viên/lớp/khóa hoặc lọc lớp.", "Scores.Index Include Student/Class/Course; phân trang 10."),
            ("6", "Admin/Staff", "Bấm Xuất PDF.", "StudentResultPdf tìm đúng Score, Include Student/Course/Class/Teacher và SimplePdfService sinh file."),
        ],
        "interaction": "Sau Teacher lưu, chuyển sang Student để xem ngay. Admin/Staff có thể tải phiếu kết quả nhưng Student hiện không có nút tự xuất PDF.",
        "constraints": [("R1", "Điểm ngoài 0..10", "Chặn cả âm và >10."), ("R2", "Nhận xét >500", "Chặn."), ("R3", "Sai lớp/Student chưa Approved", "404."), ("R4", "Trùng Score", "Unique index StudentId+CourseClassId; action upsert."), ("R5", "Không có Score khi xuất", "404 NotFound.")],
        "alternatives": [("Chưa nhập", "Teacher thấy dấu trống; Student không có thẻ cho lớp đó."), ("Cập nhật", "Ghi đè điểm/nhận xét cũ, không lưu lịch sử thay đổi."), ("Average <5", "Result='Chưa đạt'.")],
        "result": "Ghi Scores; tạo PDF kết quả; ảnh hưởng báo cáo kết quả Admin/Staff và màn hình Student.",
        "source": "ScoresController; ReportsController.StudentResultPdf; Score model; SimplePdfService",
    },
    {
        "id": "F21",
        "title": "Quản lý bài giảng và học liệu",
        "actors": "Teacher; Student; Admin; Staff; hệ thống tệp; YouTube.",
        "initiator": "Teacher/Admin/Staff thêm bài giảng; Student mở Bài giảng.",
        "entry": "CourseLectures/MyLectures/AddLecture/RemoveLecture; LearningMaterials; CRUD CourseLectures.",
        "prereq": "Teacher chỉ quản lý Course có CourseClass do mình phụ trách; Student chỉ xem Course đã Approved; Admin/Staff có quyền toàn bộ.",
        "success": "CourseLecture có tệp hoặc link YouTube; file lưu tên GUID trong wwwroot/uploads/lectures; Student mở được học liệu.",
        "limits": "Tệp <=20 MB; chỉ PDF, Word, PowerPoint, Excel, TXT, ZIP; title <=200; filename <=255; YouTube host hợp lệ.",
        "steps": [
            ("1", "Teacher", "Mở Bài giảng, chọn Course đang phụ trách, nhập title, chọn file và/hoặc YouTube.", "MyLectures chỉ cấp danh sách Course từ CourseClasses của Teacher."),
            ("2", "Hệ thống", "Kiểm tra quyền và nguồn học liệu.", "AddLecture xác minh Teacher-Course, title, ít nhất file hoặc YouTube, extension/size/host."),
            ("3", "Hệ thống", "Lưu bài giảng.", "Tệp được đổi tên GUID giữ extension; tạo CourseLecture với tên gốc, URL, YouTubeUrl, UploadedAt."),
            ("4", "Teacher", "Mở file/YouTube hoặc bấm Xóa.", "RemoveLecture chỉ xóa lecture thuộc Teacher; xóa file local an toàn rồi xóa record."),
            ("5", "Student", "Mở Bài giảng và tài liệu.", "LearningMaterials lấy CourseId từ Enrollment Approved; nạp mọi lecture thuộc các Course đó, kèm Course/Teacher."),
            ("6", "Admin/Staff", "Dùng Index/Details/Create/Edit/Delete.", "Có thể chọn bất kỳ Course/Teacher hợp lệ; Edit có thể thay file và xóa file cũ; Delete xóa record + file local."),
        ],
        "interaction": "Không có notification khi bài mới được đăng; Student phải tự mở/tải lại LearningMaterials. Link file và YouTube mở tab mới.",
        "constraints": [("R1", "Thiếu title hoặc cả file lẫn YouTube", "Không tạo."), ("R2", "Extension/size/tên file", "Chặn theo whitelist, 20 MB, 255 ký tự."), ("R3", "YouTube", "Chỉ youtube.com, subdomain youtube.com hoặc youtu.be."), ("R4", "Teacher chọn Course không phụ trách", "404."), ("R5", "Xóa file", "Chỉ xóa URL nằm dưới /uploads/lectures và đường dẫn thực nằm trong uploadRoot."), ("R6", "An toàn file", "Code kiểm tra extension nhưng chưa kiểm tra MIME/nội dung độc hại hoặc quét virus.")],
        "alternatives": [("Chỉ YouTube", "FileName='YouTube', FileUrl rỗng."), ("Chỉ file", "YouTubeUrl rỗng."), ("Edit không tải file mới", "Giữ file cũ; vẫn kiểm tra YouTube nếu nhập."), ("FileUrl ngoài upload folder", "Xóa record nhưng không xóa tài nguyên ngoài.")],
        "result": "Ghi/xóa CourseLectures và file vật lý; ảnh hưởng kho học liệu Student.",
        "source": "CourseLecturesController; CourseLecture model; related views",
    },
    {
        "id": "F22",
        "title": "Báo cáo doanh thu quý và các tài liệu PDF",
        "actors": "Admin, Staff; Student với hóa đơn của mình.",
        "initiator": "Admin/Staff chọn báo cáo/PDF; Student mở hóa đơn khi đã Paid.",
        "entry": "Reports/QuarterlyRevenue, QuarterlyRevenuePdf, InvoicePdf, StudentResultPdf.",
        "prereq": "Đúng role; dữ liệu PaymentTransaction/Payment/Score tồn tại và đạt điều kiện.",
        "success": "Hiển thị biểu đồ 3 tháng và sinh PDF doanh thu, hóa đơn hoặc phiếu kết quả.",
        "limits": "Doanh thu chỉ dùng transaction Approved có ApprovedAt; quý hợp lệ 1..4 trên trang; PDF action nhận trực tiếp year/quarter.",
        "steps": [
            ("1", "Admin/Staff", "Mở Doanh thu theo quý, chọn năm/quý và bấm Xem thống kê.", "QuarterlyRevenue chọn mặc định quý hiện tại nếu tham số không hợp lệ, gọi BuildMonthlyRevenueAsync."),
            ("2", "Hệ thống", "Tổng hợp doanh thu.", "Lọc PaymentTransaction=Approved, ApprovedAt đúng năm và 3 tháng quý; GroupBy tháng, Sum Amount; tháng không có dữ liệu=0."),
            ("3", "Admin/Staff", "Bấm Xuất PDF.", "QuarterlyRevenuePdf gọi SimplePdfService.BuildRevenueReport và tải bao-cao-doanh-thu-q{q}-{year}.pdf."),
            ("4", "Admin/Staff", "Từ Scores/Index bấm Xuất PDF kết quả.", "StudentResultPdf sinh file theo Student.Code và Class.Code."),
            ("5", "Student/Admin/Staff", "Mở hóa đơn Payment.", "InvoicePdf kiểm tra Payment tồn tại, Student sở hữu nếu role Student, Status=Paid rồi BuildInvoice."),
        ],
        "interaction": "PDF là kết quả cuối, không gửi email. Student nhận deep-link hóa đơn qua notification khi thanh toán đủ được duyệt.",
        "constraints": [("R1", "Quarter page", "Nếu year null dùng năm hiện tại; quarter ngoài 1..4 dùng quý hiện tại."), ("R2", "Invoice", "Chỉ Paid; Student không xem của người khác."), ("R3", "StudentResult", "Chỉ Admin/Staff; Score không tồn tại trả 404."), ("R4", "Revenue", "Dùng ApprovedAt chứ không dùng PaidAt/PaidDate.")],
        "alternatives": [("Quý không doanh thu", "Hiện 3 tháng bằng 0 và PDF tổng 0."), ("Payment chưa Paid", "Báo chỉ xuất hóa đơn sau khi duyệt và quay lại Payments."), ("Student mở invoice", "Trả PDF inline không đặt tên tải; Admin/Staff nhận tên file tải.")],
        "result": "Không ghi dữ liệu; xuất PDF. Doanh thu phụ thuộc tính nhất quán của PaymentTransactions ở F09/F23.",
        "source": "ReportsController; SimplePdfService; Views/Reports/QuarterlyRevenue.cshtml",
    },
    {
        "id": "F23",
        "title": "Quản trị trực tiếp lịch sử giao dịch thanh toán (CRUD)",
        "actors": "Admin, Staff.",
        "initiator": "Admin/Staff mở Lịch sử giao dịch.",
        "entry": "PaymentTransactions/Index -> Create/Edit/Details/Delete.",
        "prereq": "Payment và Student tồn tại; Student phải trùng Payment.StudentId.",
        "success": "PaymentTransaction được thêm/sửa/xóa theo form quản trị.",
        "limits": "10 giao dịch/trang; chức năng CRUD này độc lập với luồng duyệt Payments và hiện không đồng bộ lại Payment.PaidAmount/Status.",
        "steps": [
            ("1", "Admin/Staff", "Tìm theo Student, RecordedBy hoặc Note.", "Index Include Payment/Student, sắp PaidAt giảm dần, phân trang."),
            ("2", "Admin/Staff", "Bấm Thêm và chọn Payment, Student, Amount, Method, PaidAt, RecordedBy, Note.", "Create trim, TryValidateModel, kiểm tra Payment/Student và sự khớp, insert transaction."),
            ("3", "Admin/Staff", "Sửa giao dịch.", "Edit bind cùng nhóm trường rồi _context.Update; xử lý concurrency nếu id mất."),
            ("4", "Admin/Staff", "Xóa giao dịch.", "DeleteConfirmed remove và SaveChanges."),
        ],
        "interaction": "Không thông báo cho Student; không chạy ApproveTransaction/RejectTransaction.",
        "constraints": [("R1", "Amount", "0.01..1.000.000.000."), ("R2", "RecordedBy", "Bắt buộc, <=100; Note <=500."), ("R3", "Payment/Student", "Đều phải tồn tại và cùng chủ."), ("R4", "Trạng thái", "Create/Edit form không bind Status/ApprovedAt/ApprovedBy; giá trị có thể mặc định Pending khi thêm."), ("R5", "Tính nhất quán", "Thêm/sửa/xóa không cộng/trừ Payment.PaidAmount và không tính lại Payment.Status.")],
        "alternatives": [("Sai Student so với Payment", "ModelState error."), ("id sai", "404."), ("Xóa transaction Approved", "Record bị xóa nhưng Payment tổng hợp không đổi; cần xử lý bổ sung ngoài code hiện tại.")],
        "result": "Ghi PaymentTransactions; có thể ảnh hưởng báo cáo doanh thu nhưng không tự đồng bộ công nợ - đây là điểm quan trọng cần nêu khi bảo vệ.",
        "source": "PaymentTransactionsController; PaymentTransaction model; PaymentsController aggregate logic",
    },
]


PERMISSION_ROWS = [
    ("Trang chủ, khóa học, lịch khai giảng", "Xem", "Xem/CRUD khóa-lớp", "Xem/CRUD khóa-lớp", "Xem lịch của mình", "Xem/đăng ký"),
    ("Đăng ký tài khoản", "Thực hiện", "-", "-", "-", "Tạo role Student"),
    ("Duyệt đăng ký", "-", "Duyệt/hủy", "Duyệt/hủy", "-", "Gửi/hủy Pending"),
    ("Học phí", "-", "Ghi nhận/duyệt/từ chối", "Ghi nhận/duyệt/từ chối", "-", "Gửi thanh toán/xem hóa đơn"),
    ("Hồ sơ Student/Teacher", "-", "CRUD hồ sơ", "CRUD hồ sơ", "Tự sửa hồ sơ", "Tự sửa hồ sơ"),
    ("Tài khoản/vai trò", "-", "CRUD", "-", "-", "-"),
    ("Lớp/roster", "Xem lịch công khai", "CRUD/xem", "CRUD/xem", "Chỉ lớp mình", "Chỉ lịch mình"),
    ("Điểm danh", "-", "Xem tổng hợp", "Xem tổng hợp", "Nhập lớp mình", "Xem của mình"),
    ("Điểm số", "-", "Xem + PDF", "Xem + PDF", "Nhập lớp mình", "Xem của mình"),
    ("Bài giảng", "-", "CRUD toàn bộ", "CRUD toàn bộ", "CRUD phần mình", "Xem khóa Approved"),
    ("Doanh thu", "-", "Xem/PDF", "Xem/PDF", "-", "-"),
    ("Thông báo", "-", "Của mình", "Của mình", "Của mình", "Của mình"),
]


STATE_ROWS = [
    ("Enrollment", "Pending", "Student vừa đăng ký; chờ Admin/Staff."),
    ("Enrollment", "Approved", "Đã xếp lớp; tạo Payment nếu chưa có."),
    ("Enrollment", "Cancelled", "Không còn hiệu lực; CourseClassId bị gỡ; có thể đăng ký lại."),
    ("Payment", "Unpaid", "Chưa có tiền được duyệt."),
    ("Payment", "PendingApproval", "Có ít nhất một yêu cầu thanh toán đang chờ."),
    ("Payment", "PartiallyPaid", "Đã duyệt một phần, còn công nợ."),
    ("Payment", "Paid", "PaidAmount >= Amount; được xuất hóa đơn."),
    ("Payment", "Cancelled", "Khoản học phí gắn đăng ký bị hủy."),
    ("PaymentTransaction", "Pending", "Student đã khai báo thanh toán, chờ xử lý."),
    ("PaymentTransaction", "Approved", "Đã xác nhận và được tính vào doanh thu."),
    ("PaymentTransaction", "Rejected", "Không cộng PaidAmount; Student được thông báo."),
    ("CourseClass", "Upcoming/Open", "Suy ra theo StartDate nếu không Locked/Closed."),
    ("CourseClass", "Locked/Closed", "Không nhận đăng ký; Closed còn được suy ra khi đã quá EndDate."),
]


OBSERVATIONS = [
    ("Mật khẩu", "UserAccount.Password đang lưu và so sánh trực tiếp; cần băm có salt (ASP.NET Identity/PasswordHasher) trước triển khai thật.", "Cao"),
    ("OTP/đăng nhập", "Chưa có giới hạn số lần nhập, resend, captcha hoặc lockout; có nguy cơ brute-force/spam.", "Cao"),
    ("Thanh toán QR/Card/Cash", "Không có webhook/cổng thanh toán; thao tác 'Tôi đã thanh toán' chỉ tạo yêu cầu chờ duyệt.", "Trung bình"),
    ("CRUD PaymentTransaction", "Không đồng bộ Payment.PaidAmount/Status; sửa có thể làm lệch công nợ và báo cáo.", "Cao"),
    ("Điều chỉnh giảm học phí", "RecordPayment tạo transaction Approved với Amount=abs(chênh lệch); báo cáo cộng dương nên có thể làm tăng doanh thu khi giảm số đã đóng.", "Cao"),
    ("Cạnh tranh RejectTransaction", "Approve dùng Serializable, Reject không dùng transaction Serializable; hai nhân viên xử lý đồng thời có thể cần optimistic concurrency token.", "Trung bình"),
    ("Lịch lớp", "Chưa kiểm tra giáo viên/phòng trùng giờ; Schedule là chuỗi tự do ngoài việc nhận diện thứ học.", "Trung bình"),
    ("Upload bài giảng", "Có whitelist extension và giới hạn size nhưng chưa kiểm MIME, signature hay quét mã độc.", "Trung bình"),
    ("Thông báo bài giảng/đổi lịch", "Không gửi notification khi có học liệu mới hoặc thay đổi lịch lớp.", "Thấp"),
    ("Thu hồi phiên", "Tắt IsActive không xóa ngay cookie đã phát; cần security stamp/session revocation nếu yêu cầu.", "Trung bình"),
]


ORPHAN_VIEWS = [
    ("AttendanceRecords", "Create, Edit, Details, Delete", "Không có action tương ứng trong AttendanceRecordsController; chức năng chạy thực tế là Index/Manage/Save/SaveAll/MyAttendance."),
    ("Scores", "Create, Edit, Details, Delete", "Không có action tương ứng; dùng Manage/Save/MyScores/Index."),
    ("Enrollments", "Create, Edit, Delete", "Không có action CRUD; Details có, xử lý bằng Register/UpdateStatus/Cancel."),
    ("Payments", "Create, Edit, Details, Delete", "Không có action CRUD; dùng Index/RecordPayment/Pay/Approve/Reject/QR/MyPayments."),
]


def build_document() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])

    # Editorial-cover first page.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("ĐỒ ÁN NHÓM 7")
    set_run_font(r, size=11, bold=True, color=BLUE)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("KỊCH BẢN CHỨC NĂNG ĐẦY ĐỦ\nWEB QUẢN LÝ TRUNG TÂM TIẾNG ANH")
    for r in p.runs:
        set_run_font(r, size=28, bold=True, color=INK)

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Đối chiếu trực tiếp Controller - Model - Service - View - DbContext")

    add_callout(
        doc,
        "Phạm vi",
        "23 kịch bản chức năng đang có action thực thi, bao quát Khách, Admin, Staff, Teacher và Student. Nội dung mô tả đúng hành vi hiện tại của mã nguồn, kể cả ràng buộc, ngoại lệ và điểm cần lưu ý khi bảo vệ đồ án.",
        fill=LIGHTER_BLUE,
    )

    cover_meta = [
        ("Đơn vị", "Trường Đại học Kinh tế - Kỹ thuật Công nghiệp (UNETI)"),
        ("Đồ án", "Web quản lý trung tâm tiếng Anh"),
        ("Nhóm", "Nhóm 7"),
        ("Giảng viên hướng dẫn", "Phạm Thị Thùy"),
        ("Ngày lập tài liệu", "08/08/2026"),
        ("Cơ sở đối chiếu", "Mã nguồn trong workspace do_an_nhom_7"),
    ]
    add_table(doc, ["Thông tin", "Nội dung"], cover_meta, [2700, 6660], body_size=10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tài liệu dùng cho demo, thuyết trình và trả lời câu hỏi phản biện.")
    set_run_font(r, size=9.5, italic=True, color=MUTED)

    doc.add_page_break()
    doc.add_heading("1. Cách đọc và nguyên tắc phạm vi", level=1)
    add_label_paragraph(doc, "Mục tiêu", "Biến toàn bộ chức năng có thật trong code thành kịch bản có thể demo tuần tự và giải thích được dữ liệu, quyền, validate, transaction và kết quả.")
    add_label_paragraph(doc, "Quy ước", "Mỗi Fxx trả lời đủ 7 nhóm câu hỏi mẫu: tác nhân; bối cảnh; luồng chính; tương tác; validate; ngoại lệ; kết quả/bước tiếp theo.")
    add_label_paragraph(doc, "Chức năng được tính", "Action Controller có đường chạy thực tế và View/response tương ứng. View scaffold còn tồn tại nhưng không có action được tách ở phụ lục, không nhận là chức năng đang hoạt động.")
    add_label_paragraph(doc, "Anti-forgery", "Program.cs bật AutoValidateAntiforgeryToken toàn cục; các POST chính còn khai báo ValidateAntiForgeryToken tường minh.")
    add_label_paragraph(doc, "Phân trang", "Courses dùng 6/trang; phần lớn danh sách quản trị dùng 10/trang; Notifications lấy tối đa 50 mục mới nhất.")

    doc.add_heading("2. Ma trận quyền theo vai trò", level=1)
    add_table(
        doc,
        ["Nhóm chức năng", "Khách", "Admin", "Staff", "Teacher", "Student"],
        PERMISSION_ROWS,
        [1950, 1050, 1700, 1700, 1450, 1510],
        body_size=8.5,
    )
    add_callout(doc, "Lưu ý", "Một số action gắn [AllowAnonymous] nhưng tự kiểm tra User.IsInRole và claim bên trong (Profile, MySavedCourses, MyLectures, LearningMaterials). Kết quả bảo vệ quyền vẫn là Challenge/Forbid/NotFound tùy trường hợp.")

    doc.add_heading("3. Chuỗi trạng thái nghiệp vụ", level=1)
    add_table(doc, ["Đối tượng", "Trạng thái", "Ý nghĩa trong hệ thống"], STATE_ROWS, [2000, 1900, 5460], body_size=9.2)

    doc.add_heading("4. Danh mục kịch bản", level=1)
    index_rows = [(s["id"], s["title"], s["actors"]) for s in SCENARIOS]
    add_table(doc, ["Mã", "Chức năng", "Tác nhân chính"], index_rows, [750, 3860, 4750], body_size=9.0)

    for position, scenario in enumerate(SCENARIOS, start=1):
        doc.add_page_break()
        doc.add_heading(f"{scenario['id']} - {scenario['title']}", level=1)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f"Kịch bản {position}/{len(SCENARIOS)}")
        set_run_font(r, size=9, bold=True, color=MUTED)

        doc.add_heading("Nhóm 1-2. Tác nhân, bối cảnh và mục tiêu", level=2)
        add_table(doc, ["Thuộc tính", "Nội dung"], metadata_rows(scenario), [2700, 6660], body_size=9.6)

        doc.add_heading("Nhóm 3. Luồng chính", level=2)
        add_table(doc, ["Bước", "Người thực hiện", "Thao tác", "Điều xảy ra trong hệ thống"], scenario["steps"], [620, 1450, 2600, 4690], body_size=8.9)

        doc.add_heading("Nhóm 4. Tương tác và điểm chuyển khi demo", level=2)
        add_label_paragraph(doc, "Kịch bản tương tác", scenario["interaction"])

        doc.add_heading("Nhóm 5. Ràng buộc và validate", level=2)
        add_table(doc, ["Mã", "Điều kiện", "Phản ứng hệ thống"], scenario["constraints"], [700, 3650, 5010], body_size=9.0)

        doc.add_heading("Nhóm 6. Luồng phụ và ngoại lệ", level=2)
        add_table(doc, ["Trường hợp", "Kết quả/đường đi tiếp"], scenario["alternatives"], [3300, 6060], body_size=9.2)

        doc.add_heading("Nhóm 7. Kết quả và bước tiếp theo", level=2)
        result_paragraph = add_label_paragraph(doc, "Dữ liệu/ảnh hưởng", scenario["result"])
        result_paragraph.paragraph_format.keep_with_next = True
        add_label_paragraph(doc, "Đối chiếu code", scenario["source"], color=DARK_BLUE, after=0)

    doc.add_page_break()
    doc.add_heading("Phụ lục A. Kịch bản demo liên vai trò đề xuất", level=1)
    demo_steps = [
        ("1", "Trình duyệt A - Student", "Đăng nhập, mở chi tiết khóa, chọn lớp còn chỗ và đăng ký."),
        ("2", "Trình duyệt B - Admin/Staff", "Quan sát badge; mở thông báo/Enrollments; chọn lớp và duyệt."),
        ("3", "Trình duyệt A - Student", "Mở thông báo duyệt, đi thẳng MyPayments; chọn chuyển khoản và mở QR."),
        ("4", "Trình duyệt A - Student", "Bấm Tôi đã thanh toán để tạo PaymentTransaction Pending."),
        ("5", "Trình duyệt B - Admin/Staff", "Mở Payments; duyệt transaction; kiểm tra Dashboard và báo cáo."),
        ("6", "Trình duyệt A - Student", "Mở thông báo, xem trạng thái Paid và hóa đơn PDF."),
        ("7", "Trình duyệt C - Teacher", "Mở lớp giảng dạy; kiểm tra roster; điểm danh, nhập điểm và đăng bài giảng."),
        ("8", "Trình duyệt A - Student", "Tải lại lịch học, điểm danh, kết quả và bài giảng để thấy dữ liệu Teacher vừa tạo."),
        ("9", "Trình duyệt B - Admin/Staff", "Mở lịch sử điểm danh, kết quả và xuất PDF kết quả/doanh thu."),
    ]
    add_table(doc, ["Bước", "Màn hình", "Thao tác/điểm quan sát"], demo_steps, [650, 2400, 6310], body_size=9.4)

    doc.add_heading("Phụ lục B. Điểm cần chủ động trả lời khi bảo vệ", level=1)
    add_table(doc, ["Chủ đề", "Hiện trạng trong code", "Mức độ"], OBSERVATIONS, [2100, 6260, 1000], body_size=9.0)
    add_callout(doc, "Cách trình bày", "Nêu rõ đâu là kiểm soát đã có (Authorize, antiforgery, unique index, Serializable transaction, giới hạn upload) và đâu là cải tiến cho môi trường thật. Không nên khẳng định hệ thống đã có cổng thanh toán tự động hoặc mã hóa mật khẩu vì code hiện chưa có.", fill="FFF8E8", color=GOLD)

    doc.add_heading("Phụ lục C. View tồn tại nhưng chưa phải chức năng chạy", level=1)
    add_table(doc, ["Nhóm", "View scaffold", "Kết luận"], ORPHAN_VIEWS, [1900, 2500, 4960], body_size=9.2)

    doc.add_heading("Phụ lục D. Bảng dữ liệu và ràng buộc trọng yếu", level=1)
    data_rows = [
        ("Roles", "Name unique", "Phân quyền; không xóa khi UserAccount đang dùng."),
        ("UserAccounts", "UserName unique; StudentId/TeacherId unique khi có", "Một hồ sơ tối đa một tài khoản."),
        ("Students", "Code, Email unique", "Tuổi 5-100; dữ liệu học tập phần lớn Restrict."),
        ("Teachers", "Code, Email unique", "CourseClass và CourseLecture Restrict."),
        ("Courses", "Code unique", "Tuition precision 18,2."),
        ("CourseClasses", "Code unique; index StartDate", "Capacity 1-500; status lưu chuỗi."),
        ("Enrollments", "StudentId+CourseId unique khi Status != Cancelled", "Cho đăng ký lại sau hủy."),
        ("Payments", "EnrollmentId unique", "Một Enrollment tối đa một Payment."),
        ("Scores", "StudentId+CourseClassId unique", "Một kết quả/lớp/học viên."),
        ("AttendanceRecords", "StudentId+CourseClassId+StudyDate unique", "Một điểm danh/buổi/học viên."),
        ("SavedCourses", "StudentId+CourseId unique", "Không lưu trùng."),
        ("Notifications", "Index UserAccountId+IsRead+CreatedAt", "Tối ưu badge/danh sách."),
    ]
    add_table(doc, ["Bảng", "Chỉ mục/ràng buộc", "Ý nghĩa nghiệp vụ"], data_rows, [1900, 3370, 4090], body_size=9.0)

    doc.add_heading("Kết luận", level=1)
    add_label_paragraph(doc, "Phạm vi hoàn tất", f"Tài liệu đã bao phủ {len(SCENARIOS)} nhóm chức năng có action chạy thực tế, từ khách truy cập đến quy trình liên vai trò và quản trị nội bộ.")
    add_label_paragraph(doc, "Chuỗi nghiệp vụ trung tâm", "Course -> CourseClass -> Enrollment -> Payment -> PaymentTransaction -> Invoice/Revenue; song song là CourseClass -> Attendance/Score và Course -> CourseLecture.")
    add_label_paragraph(doc, "Thông điệp bảo vệ", "Hệ thống đã có phân quyền theo role, antiforgery, kiểm tra sở hữu, unique index và transaction Serializable ở các điểm tranh chấp chính. Các quan sát trong Phụ lục B là phạm vi cải tiến, không phải chức năng đã hoàn thành.")

    doc.core_properties.title = "Kịch bản chức năng Web quản lý Trung tâm Tiếng Anh - Nhóm 7"
    doc.core_properties.subject = "Kịch bản chi tiết theo 7 nhóm câu hỏi, đối chiếu mã nguồn"
    doc.core_properties.author = "Nhóm 7"
    doc.core_properties.keywords = "ASP.NET Core MVC, use case, kịch bản chức năng, trung tâm tiếng Anh"
    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_document())
